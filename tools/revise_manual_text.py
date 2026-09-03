from pathlib import Path
from tempfile import TemporaryDirectory
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document


SOURCE = Path(r"C:\xampp\htdocs\wordpress\wp-content\plugins\Obatala\Manual de Curadoria 3.1 - atualizado.docx")
OUTPUT = Path(r"C:\xampp\htdocs\wordpress\wp-content\plugins\Obatala\Manual de Curadoria 3.1 - revisado.docx")


REPLACEMENTS = {
    12: "Este manual orienta o uso do Obatalá e reúne informações, procedimentos e exemplos para apoiar a implantação e a operação do sistema.",
    13: "As orientações aqui apresentadas devem ser aplicadas de acordo com a realidade de cada instituição. Em caso de dúvida, consulte a equipe técnica responsável pelo ambiente ou o suporte especializado.",
    14: "O conteúdo poderá ser atualizado para acompanhar mudanças no sistema, ajustes de interface e melhorias de uso.",
    18: "O Obatalá é um sistema de gestão curatorial desenvolvido como plugin para o WordPress, integrado ao Tainacan. Ele apoia a organização de processos ligados ao ciclo curatorial, como patrimoniação, conservação, empréstimo, circulação de itens e acompanhamento de atividades do acervo.",
    19: "O projeto é financiado pelo Instituto Brasileiro de Museus (IBRAM) e desenvolvido pelo NOCS-Lab (Laboratório Multiusuário de Pesquisa em Redes e Sistemas Computacionais) do IFRN Parnamirim. A proposta se inspira em práticas do Museu Paulista da USP e adapta esses fluxos ao ambiente WordPress/Tainacan.",
    20: "Na prática, o sistema ajuda a transformar rotinas curatoriais em processos estruturados, com etapas, responsáveis, formulários, documentos e registros de acompanhamento.",
    21: "Como o Obatalá segue em desenvolvimento, a participação dos usuários é parte importante do processo de melhoria. O ambiente de testes permite experimentar funcionalidades, registrar dificuldades, sugerir ajustes e validar se os fluxos atendem às necessidades reais das instituições.",
    23: "Plugin do WordPress",
    24: "Por ser um plugin do WordPress, o Obatalá precisa de uma instalação do WordPress em uma máquina local ou em um servidor da instituição. O Tainacan também deve estar instalado e configurado, pois é a base usada para gestão e publicação dos acervos. Para orientações sobre a instalação do Tainacan, consulte os vídeos oficiais do canal do projeto no YouTube: https://www.youtube.com/watch?v=7v6qNHmqm0I e https://www.youtube.com/watch?v=qRtoNRUlVkk&t=15s.",
    25: "Depois de configurar WordPress e Tainacan, a equipe técnica pode instalar e ativar o plugin Obatalá. Recomenda-se seguir as boas práticas de segurança, atualização e desempenho adotadas pela instituição.",
    26: "Login no WordPress",
    27: "O acesso ao Obatalá usa as mesmas credenciais do painel administrativo do WordPress. O nome de usuário e a senha são definidos durante a configuração inicial ou criados posteriormente pelo administrador do ambiente.",
    28: "Na tela de login, informe o nome de usuário e a senha cadastrados. Caso não tenha esses dados, solicite apoio ao administrador do sistema no museu.",
    31: "Nota: mantenha nome de usuário e senha em local seguro. Se houver dificuldade de acesso, use a opção de recuperação de senha da tela de login ou procure o responsável técnico pelo ambiente.",
    33: "Menu administrativo do WordPress",
    34: "Após a instalação, o Obatalá aparece no menu lateral do painel administrativo do WordPress. A opção “Obatalá” dá acesso às áreas do sistema voltadas à gestão curatorial.",
    37: "Ao clicar em Obatalá, o sistema abre o painel inicial. Essa tela apresenta indicadores de processos, dados do usuário, grupo associado, quantidade de itens do acervo, modelos, grupos, notificações e modelos mais usados.",
    40: "O menu superior permanece visível nas telas administrativas do plugin. Ele reúne atalhos para Dashboard, Processos, Modelos, Itens do acervo, Tainacan, Configurações e WordPress. Em Configurações, o acesso a Grupos e Usuários depende das permissões da conta conectada.",
    42: "Dashboard: exibe uma visão geral do ambiente, com contadores, progresso dos processos, modelos mais usados e notificações.",
    43: "Modelos: permite criar e gerenciar modelos de processo, além de configurar etapas, campos, condicionais e mapeamentos para exportação ao Tainacan.",
    44: "Processos: permite criar, acompanhar, preencher, comentar, gerar relatórios e concluir processos curatoriais baseados nos modelos cadastrados.",
    45: "As próximas seções detalham essas funcionalidades e mostram como utilizá-las na rotina de gestão do acervo.",
    49: "A Gerência de Grupos organiza os usuários do museu em equipes ou setores. Essa organização ajuda a distribuir responsabilidades e controlar quem pode atuar em cada processo.",
    50: "Nessa área, é possível criar grupos, alterar nomes e descrições e administrar os usuários associados.",
    51: "Para criar um grupo, clique em Adicionar novo. Informe título e descrição na janela exibida e salve. O grupo passa a aparecer na listagem com status, quantidade de usuários e ações de visualização ou gerenciamento.",
    55: "Depois de cadastrado, o grupo aparece na lista de grupos do sistema.",
    58: "Na lista de grupos, use as ações disponíveis para visualizar detalhes, editar informações, gerenciar usuários ou excluir o grupo quando necessário.",
    65: "Ao clicar em Gerenciar usuários ou Ver grupo, o sistema abre os detalhes do grupo e a lista de usuários associados. Nessa área, é possível conferir nome, nome de usuário, e-mail e remover usuários, desde que a conta tenha permissão para isso.",
    68: "Somente usuários já cadastrados no WordPress podem ser adicionados a um grupo. Se o usuário ainda não existir, acesse o menu lateral do WordPress, entre em Users e clique em Add New User. Em seguida, preencha os dados solicitados.",
    71: "Para remover um usuário de um grupo, clique em Manage users, localize o usuário na janela exibida e use o ícone de lixeira correspondente. A remoção vale apenas para aquele grupo.",
    75: "A área Models, disponível no menu superior, permite criar, visualizar e configurar modelos de processos curatoriais. Esses modelos definem etapas e fluxos para atividades como patrimoniação, conservação, empréstimo e circulação de itens de acervo.",
    76: "Nesta seção, você verá como criar e ajustar modelos de processo para a realidade da instituição. Um modelo bem configurado padroniza a execução, facilita o acompanhamento e reduz a perda de informações ao longo do fluxo.",
    77: "A equipe também estuda a adaptação da notação BPMN (Business Process Model and Notation) ao desenho dos fluxos no Obatalá. A intenção é aproveitar conceitos já consolidados para tornar os processos mais claros para equipes técnicas e curatoriais.",
    78: "Para adicionar um modelo, clique em Adicionar modelo de processo. O cadastro inicial solicita título e descrição. Depois de salvo, o modelo pode ser aberto pela ação Gerenciar modelo, que concentra etapas, campos, condicionais e configurações de exportação para o Tainacan.",
    83: "Após o cadastro, o modelo pode ser editado ou excluído pelas ações disponíveis em sua linha na listagem. Use essas opções para manter os modelos atualizados conforme os fluxos da instituição mudarem.",
    87: "A principal ação da listagem é Gerenciar modelo. Essa tela reúne o editor visual do fluxo e o acordeon Exportação para o Tainacan. Usuários autorizados podem definir etapas, campos, conexões, condicionais e o relacionamento entre campos do Obatalá e metadados do Tainacan.",
    90: "A tela de edição oferece controles para salvar alterações, cancelar mudanças, adicionar etapa, adicionar condicional, alternar tela cheia, importar/exportar JSON e abrir o acordeon Exportação para o Tainacan. Esse acordeon aparece acima de Gerenciar etapas, indica se o mapeador está Ativo ou Inativo e reúne a habilitação da exportação e a seleção das coleções de destino.",
    99: "Exportar JSON: exporta o modelo de processo em formato JSON, útil para armazenamento, backup ou compartilhamento com outra instituição.",
    100: "Importar JSON: importa um modelo previamente exportado para edição ou reutilização.",
    101: "Essa área permite ajustar o modelo de processo de acordo com as rotinas do museu. O editor visual facilita a organização das etapas e a conferência do fluxo antes de colocá-lo em uso.",
    102: "Para iniciar a edição de um modelo, clique em Add step. Uma nova etapa será criada na área de edição, representada por um retângulo com nome padrão e botões de configuração.",
    103: "Em um modelo de Restauração, por exemplo, a primeira atividade pode ser Avaliação do item a ser restaurado. Essa etapa pode reunir critérios de avaliação, responsáveis, prazos e observações iniciais.",
    105: "A imagem a seguir destaca os principais recursos de uma etapa: exclusão da etapa, criação de campos do formulário e botões de conexão usados para desenhar o fluxo entre etapas.",
    114: "O botão de adição de campos permite incluir tipos como Texto, Email, Número, Seletor de data, Upload de arquivo, Documento de etapa, Select, Radio e Busca no Tainacan. Quando a exportação para o Tainacan está ativa, os campos também podem exibir os controles Coleção de destino e Metadado de destino, usados para mapear o dado do Obatalá ao metadado correspondente no Tainacan.",
    119: "Importante: cada campo possui propriedades próprias de configuração. Essas propriedades ajudam a orientar o preenchimento e a controlar como a informação será registrada no formulário.",
    120: "Como exemplo, considere a inclusão de um campo Texto na etapa Testes Preliminares. Esse tipo de campo registra informações descritivas e pode ser configurado com as seguintes propriedades:",
    121: "Label: nome ou título exibido para identificar o campo.",
    122: "Placeholder: texto exibido dentro do campo antes do preenchimento, usado como orientação ou exemplo.",
    123: "Obrigatório: define se o campo precisa ser preenchido para avançar na etapa.",
    124: "Tamanho mínimo e máximo: define limites de caracteres para o valor informado.",
    125: "Padrão de avaliação: define critérios de validação, como formato ou conteúdo esperado.",
    126: "Texto de ajuda: orientação adicional para apoiar o preenchimento correto.",
    129: "Essas propriedades tornam o campo Texto adaptável a diferentes usos. No exemplo, ele pode registrar observações dos testes preliminares. Em seguida, o usuário pode adicionar um campo de data, com um conjunto menor de propriedades:",
    130: "Label: nome ou título exibido para identificar o campo.",
    131: "Obrigatório: define se o campo precisa ser preenchido para avançar na etapa.",
    132: "Texto de ajuda: orientação adicional para apoiar o preenchimento correto.",
    136: "Com exceção da Busca no Tainacan, os demais tipos de campo possuem propriedades configuráveis. A imagem a seguir mostra opções do campo Upload de arquivo.",
    139: "Cada etapa pode ser conectada a outras etapas para formar o fluxo completo do processo. A sequência deve refletir a ordem real das atividades curatoriais e pode ser ajustada sempre que o modelo precisar ser revisto.",
    142: "Importante: as etapas são conectadas por fluxos criados a partir do botão azul. Use essas conexões para ligar uma etapa à próxima e manter uma sequência lógica, preferencialmente organizada da esquerda para a direita.",
    144: "Importante: depois de editar o modelo, clique em Save para gravar as alterações.",
    146: "O elemento condicional cria ramificações no fluxo. Ao configurar uma condicional, o usuário seleciona um campo do formulário e informa um valor de referência. Durante a execução, a resposta registrada nesse campo define qual caminho o processo deve seguir.",
    150: "A modelagem de processos define a estrutura que será seguida durante a execução das atividades curatoriais. Cada etapa modelada se transforma em um formulário no sistema, permitindo registrar informações de modo padronizado e acompanhar o andamento do trabalho.",
    152: "Depois de criar um modelo, o próximo passo é utilizá-lo para abrir um processo. O processo é a execução prática do modelo: ele reúne etapas, formulários, responsáveis e registros produzidos durante o fluxo.",
    153: "Na área Processos, é possível criar processos a partir dos modelos cadastrados. A listagem oferece abas para Todos os processos e Meus processos, filtros por nível de acesso, modelo e progresso, além de ações para visualizar, editar, consultar histórico, excluir e gerar relatório PDF quando disponível.",
    154: "Essa área permite acompanhar processos em andamento e verificar se as etapas foram preenchidas, salvas, submetidas ou concluídas. Também ajuda gestores a identificar pendências e acompanhar a rotina operacional.",
    155: "Para criar um processo, acesse Processos no menu superior do Obatalá e clique em Adicionar novo. Informe o título, selecione o modelo e defina o nível de acesso: Não restrito ou Restrito. Após salvar, o processo aparece na listagem e pode ser executado.",
    160: "Após criar o processo, clique em View, representado pelo ícone de olho, para abrir a tela de execução. Nessa área, o usuário preenche os formulários das etapas e acompanha o progresso do processo.",
    165: "Cada formulário corresponde a uma etapa do processo. O usuário pode registrar informações, salvar rascunhos, submeter etapas, adicionar comentários e consultar o histórico. Quando o modelo possui exportação Tainacan ativa, o processo também pode exibir uma etapa virtual de preparação da exportação, usada para definir coleção de destino, quantidade de itens, entrada manual ou por planilha e regras de repetição de valores.",
    168: "Depois de submeter um formulário, o usuário pode avançar para a próxima etapa. Etapas concluídas aparecem em verde, facilitando a leitura do progresso do processo.",
    171: "A área Processos transforma os modelos cadastrados em fluxos de trabalho executáveis. Com ela, os usuários iniciam, acompanham e concluem processos, mantendo histórico, registros e rastreabilidade das atividades realizadas.",
    174: "Esta seção complementa as orientações anteriores com base na interface administrativa atual do Obatalá. As telas foram verificadas no ambiente WordPress local em 31/08/2026 e refletem a organização recente do menu superior e dos recursos integrados ao Tainacan.",
    176: "O menu superior organiza o uso diário do sistema em Dashboard, Processos, Modelos, Itens do acervo, Tainacan, Configurações e WordPress. O Dashboard reúne indicadores gerais e funciona como ponto de partida para acompanhar a atividade do ambiente.",
    187: "A tela Processos mostra o total de processos, as abas Todos os processos e Meus processos, filtros e a ação Adicionar novo. Quando não há processos cadastrados, o sistema exibe uma mensagem de estado vazio.",
    190: "Ao adicionar um processo, informe o título, escolha o modelo de processo e selecione o nível de acesso. O nível Não restrito permite acesso mais amplo, conforme as permissões do usuário. O nível Restrito limita a operação aos usuários autorizados pelo fluxo e pelos grupos associados.",
    194: "A listagem de Modelos permite criar um modelo com título e descrição. Depois do cadastro, a ação principal é Gerenciar modelo. A antiga separação entre edição de etapas e edição dos dados de exportação foi unificada: o editor do modelo agora concentra fluxo, campos das etapas e acordeon Exportação para o Tainacan.",
    199: "No editor do modelo, a configuração de exportação aparece no acordeon Exportação para o Tainacan, acima de Gerenciar etapas. Quando está recolhido, o acordeon informa se o recurso está Ativo ou Inativo. Ao expandi-lo, o usuário encontra o controle Exporter status; ao ativá-lo, o sistema exibe a área Coleções de destino e o seletor de coleções.",
    200: "Com a exportação habilitada e ao menos uma coleção selecionada, os componentes de cada etapa passam a exibir uma propriedade adicional de mapeamento. Essa propriedade aparece no bloco Tainacan Export das configurações do campo e liga o campo do Obatalá a um metadado do Tainacan. Assim, o dado preenchido no processo já fica preparado para compor o item exportado.",
    201: "Exporter status: ativa ou desativa o mapeamento para exportação.",
    202: "Coleções de destino: define quais coleções do Tainacan podem receber itens gerados pelo processo.",
    203: "Aviso de alerta: aparece quando a exportação está ativa, mas nenhuma coleção foi selecionada.",
    204: "Target collection: define a coleção de destino do campo.",
    205: "Target metadata: define o metadado do Tainacan que receberá o valor do campo.",
    206: "Do not map this field: mantém o campo apenas no formulário do Obatalá, sem envio para o Tainacan.",
    207: "Resumo do mapeamento: mostra, por coleção, a relação entre Campo Obatalá e Metadado Tainacan.",
    211: "A área Itens do acervo aproxima os processos curatoriais dos registros mantidos no Tainacan. A listagem permite buscar e filtrar itens, visualizar coleção, situação, processos vinculados, última atualização e ações de consulta.",
    214: "Ao abrir um item, a tela apresenta dados do registro, metadados, botões para abrir ou editar o item no Tainacan e a linha do tempo de processos vinculados. Essa visão ajuda a identificar se o item já participou de processos curatoriais e qual é a situação desses processos.",
    218: "A área Grupos lista os grupos cadastrados, com status, descrição, quantidade de usuários e ações disponíveis. O detalhe do grupo mostra os usuários associados e permite administrar sua composição conforme as permissões da conta conectada.",
    224: "Quando um modelo possui mapeamento Tainacan ativo, cada processo criado a partir dele apresenta uma etapa de controle chamada Preparação da exportação para o Tainacan. Nessa etapa, o usuário informa como os itens serão gerados: manualmente, preenchendo os dados no próprio processo, ou via planilha, enviando um arquivo estruturado com os valores dos metadados mapeados.",
    225: "1. Selecionar a coleção de destino, quando houver mais de uma coleção mapeada.",
    226: "2. Definir se o processo gerará um item ou múltiplos itens no Tainacan.",
    227: "3. Para múltiplos itens, informar a quantidade e escolher a fonte dos dados: entrada manual ou planilha.",
    228: "4. Na entrada manual, preencher os valores diretamente nos campos das etapas.",
    229: "5. Na entrada por planilha, enviar um arquivo CSV, XLS ou XLSX compatível com os campos mapeados.",
    230: "6. Opcionalmente, repetir valores-base entre itens e definir um campo identificador único com prefixo.",
    231: "7. Salvar a preparação da exportação antes de concluir a etapa de exportação.",
    232: "A exportação cria ou atualiza itens no Tainacan com base nos metadados mapeados. O sistema também registra a referência do processo nos itens vinculados, o que facilita a rastreabilidade entre o processo curatorial e o acervo.",
    234: "A listagem de processos pode exibir a ação Gerar relatório PDF. O relatório reúne dados do processo, número do processo, etapa atual, nível de acesso e informações registradas ao longo do fluxo. Campos do tipo Documento de etapa também podem gerar PDF e exigir anexação de arquivo assinado quando essa opção estiver configurada no modelo.",
    236: "A interface exibe menus e controles conforme as permissões da conta conectada. Administradores possuem acesso amplo; gestores de modelos configuram modelos; gestores de grupos administram grupos e usuários; gestores de mapeadores acessam os controles Tainacan Export; usuários de setor interagem com os processos e etapas aos quais têm acesso. Quando a conta não possui permissão, a ação pode não aparecer ou retornar uma mensagem de acesso negado.",
}


def replace_paragraph_text(paragraph, new_text):
    paragraph_element = paragraph._p
    paragraph_properties = paragraph_element.pPr
    for child in list(paragraph_element):
        if child is not paragraph_properties:
            paragraph_element.remove(child)
    paragraph.add_run(new_text)


def main():
    doc = Document(SOURCE)

    for idx, new_text in REPLACEMENTS.items():
        if idx >= len(doc.paragraphs):
            raise IndexError(f"Paragraph {idx} does not exist")
        replace_paragraph_text(doc.paragraphs[idx], new_text)

    doc.save(OUTPUT)
    replace_ooxml_text(OUTPUT, {"Wordpress": "WordPress"})
    print(OUTPUT)


def replace_ooxml_text(docx_path, replacements):
    with TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        with ZipFile(docx_path, "r") as zin:
            zin.extractall(tmp_path)
        for xml_path in (tmp_path / "word").glob("*.xml"):
            text = xml_path.read_text(encoding="utf-8")
            updated = text
            for old, new in replacements.items():
                updated = updated.replace(old, new)
            if updated != text:
                xml_path.write_text(updated, encoding="utf-8")
        tmp_docx = docx_path.with_suffix(".tmp.docx")
        with ZipFile(tmp_docx, "w", ZIP_DEFLATED) as zout:
            for file_path in tmp_path.rglob("*"):
                if file_path.is_file():
                    zout.write(file_path, file_path.relative_to(tmp_path).as_posix())
        tmp_docx.replace(docx_path)


if __name__ == "__main__":
    main()
