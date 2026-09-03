from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


DOCX = Path(r"C:\xampp\htdocs\wordpress\wp-content\plugins\Obatala\Manual de Curadoria 3.1 - atualizado.docx")
SCREEN_DIR = Path(r"C:\xampp\htdocs\wordpress\wp-content\plugins\Obatala\manual-update-screenshots")


def set_paragraph_text(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text


def add_body_paragraph(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.style = doc.styles["normal"]
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08
    return paragraph


def add_caption(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(10)
    for run in paragraph.runs:
        run.italic = True
        run.font.size = Pt(9)
    return paragraph


def add_screenshot(doc, filename, caption):
    image_path = SCREEN_DIR / filename
    if image_path.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(6.4))
        add_caption(doc, caption)


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Paragraph")
        paragraph.add_run("- " + item)


def add_numbered(doc, items):
    for index, item in enumerate(items, start=1):
        paragraph = doc.add_paragraph(style="List Paragraph")
        paragraph.add_run(f"{index}. {item}")


doc = Document(DOCX)

# Light-touch updates to existing text.
replacements = {
    "Guia do usuário: Obatalá — Um sistema de gestão": "Guia do usuário: Obatalá - Sistema de gestão",
    "curatorial de museus": "curatorial de museus",
}

for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    if text in replacements:
        set_paragraph_text(paragraph, replacements[text])

    if text.startswith("Ao clicar no menu Obatalá, será exibido o dashboard"):
        set_paragraph_text(
            paragraph,
            "Ao clicar no menu Obatalá, será exibido o painel do sistema. A tela inicial apresenta indicadores de processos concluídos, saudação ao usuário, grupo associado, quantidade de itens do acervo, processos, modelos e grupos, além de áreas para modelos mais usados e notificações. O painel funciona como ponto de entrada para as rotinas de gestão curatorial.",
        )

    if text.startswith("Diversos elementos da página, especialmente os cards"):
        set_paragraph_text(
            paragraph,
            "O menu superior do Obatalá permanece visível nas telas administrativas do plugin. Ele reúne atalhos para Dashboard, Processos, Modelos, Itens do acervo, Tainacan, Configurações e WordPress. Em Configurações, o usuário acessa Grupos e Usuários, conforme as permissões disponíveis para sua conta.",
        )

    if text.startswith("Groups: Permite"):
        set_paragraph_text(
            paragraph,
            "Dashboard: apresenta a visão geral do ambiente, com contadores, progresso dos processos, modelos mais usados e notificações.",
        )

    if text.startswith("Models: Permite"):
        set_paragraph_text(
            paragraph,
            "Modelos: permite criar e gerenciar modelos de processo, configurar etapas, campos, condicionais e o mapeamento de exportação para o Tainacan.",
        )

    if text.startswith("Processes: Oferece"):
        set_paragraph_text(
            paragraph,
            "Processos: permite criar, acompanhar, preencher, comentar, gerar relatórios e concluir processos curatoriais baseados nos modelos cadastrados.",
        )

    if text.startswith("Para inserir um novo grupo"):
        set_paragraph_text(
            paragraph,
            "Para inserir um novo grupo, o usuário deve clicar no botão Adicionar novo. Em seguida, uma janela suspensa solicita o título e a descrição do grupo. Depois de salvo, o grupo passa a aparecer na listagem com status, quantidade de usuários e ações de visualização ou gerenciamento de usuários.",
        )

    if text.startswith("Ao clicar na opção \"Manage users\""):
        set_paragraph_text(
            paragraph,
            "Ao clicar na opção Gerenciar usuários ou Ver grupo, o usuário acessa os detalhes do grupo e a lista de usuários associados. Nessa área é possível conferir nome, nome de usuário, e-mail e remover usuários do grupo quando a conta possuir permissão para essa operação.",
        )

    if text.startswith("Para adicionar novos modelos de processos"):
        set_paragraph_text(
            paragraph,
            "Para adicionar novos modelos de processos, o usuário deve clicar no botão Adicionar modelo de processo. O cadastro inicial solicita título e descrição. Depois de salvo, o modelo pode ser aberto pela ação Gerenciar modelo, que concentra a edição das etapas, dos campos, das condicionais e das configurações de exportação para o Tainacan.",
        )

    if text.startswith("No que diz respeito à edição de um modelo"):
        set_paragraph_text(
            paragraph,
            "No que diz respeito à edição de um modelo de processo, a ação principal da listagem é Gerenciar modelo. Essa tela reúne o editor visual do fluxo e o acordeon Exportação para o Tainacan. O usuário pode configurar etapas, campos, conexões, condicionais e, quando autorizado, definir como os campos do Obatalá serão relacionados aos metadados das coleções do Tainacan.",
        )

    if text.startswith("A interface de edição é intuitiva"):
        set_paragraph_text(
            paragraph,
            "A interface de edição conta com controles para salvar alterações, cancelar mudanças, adicionar etapa, adicionar condicional, alternar tela cheia, importar/exportar JSON e abrir o acordeon Exportação para o Tainacan. Esse acordeon fica acima de Gerenciar etapas, indica se o mapeador está Ativo ou Inativo e concentra a habilitação da exportação e a seleção das coleções de destino.",
        )

    if text.startswith("Save: Salva"):
        set_paragraph_text(paragraph, "Salvar alterações: grava o fluxo do modelo e, quando configurado, as informações de mapeamento da exportação para o Tainacan.")

    if text.startswith("Cancel changes"):
        set_paragraph_text(paragraph, "Cancelar alterações: descarta mudanças feitas desde a última gravação.")

    if text.startswith("Add step"):
        set_paragraph_text(paragraph, "Adicionar etapa: cria uma nova atividade do processo e abre a configuração de campos da etapa.")

    if text.startswith("Add conditional"):
        set_paragraph_text(paragraph, "Adicionar condicional: cria uma ramificação do fluxo com base no valor informado em um campo da etapa.")

    if text.startswith("Fullscreen"):
        set_paragraph_text(paragraph, "Tela cheia: amplia a área de edição visual do fluxo.")

    if text.startswith("O botão para adição de campos"):
        set_paragraph_text(
            paragraph,
            "O botão para adição de campos permite selecionar tipos como Texto, Email, Número, Seletor de data, Upload de arquivo, Documento de etapa, Select, Radio e Busca no Tainacan. Quando a exportação para o Tainacan está ativa, os formulários de campo também podem exibir os controles Coleção de destino e Metadado de destino, usados para mapear o dado do Obatalá ao metadado correspondente no Tainacan.",
        )

    if text.startswith("No Processes, é possível criar novos processos"):
        set_paragraph_text(
            paragraph,
            "Em Processos, é possível criar processos a partir dos modelos previamente configurados. A listagem oferece abas para Todos os processos e Meus processos, filtros por nível de acesso, modelo e progresso, além de ações para visualizar, editar, consultar histórico, excluir e gerar relatório PDF quando disponível.",
        )

    if text.startswith("Para criar um novo processo"):
        set_paragraph_text(
            paragraph,
            "Para criar um novo processo, acesse Processos no menu superior do Obatalá e clique em Adicionar novo. Informe o título do processo, selecione o modelo de processo e defina o nível de acesso: Não restrito ou Restrito. Após salvar, o processo passa a ser acompanhado pela listagem e pela tela de execução das etapas.",
        )

    if text.startswith("Cada formulário exibido corresponde"):
        set_paragraph_text(
            paragraph,
            "Cada formulário exibido corresponde a uma etapa do processo. O usuário registra informações, salva rascunhos, submete etapas, adiciona comentários e acompanha o histórico. Quando o modelo possui exportação Tainacan ativa, o processo pode exibir uma etapa virtual de Preparação da exportação para definir coleção de destino, quantidade de itens, entrada manual ou por planilha e regras para repetição de valores.",
        )

# Revision history.
if doc.tables:
    table = doc.tables[0]
    existing_versions = {row.cells[0].text.strip() for row in table.rows}
    if "3.1" not in existing_versions:
        row = table.add_row()
        row.cells[0].text = "3.1"
        row.cells[1].text = "31/08/2026"
        row.cells[2].text = "Nocs lab"
        row.cells[3].text = "Atualização da interface, integração com itens do acervo, exportação Tainacan, relatórios e permissões."

# Add updated interface section.
doc.add_page_break()
doc.add_heading("Atualização da interface do Obatalá", level=1)
add_body_paragraph(
    doc,
    "Esta seção complementa as orientações anteriores com base na interface administrativa atual do Obatalá. As telas foram verificadas no ambiente WordPress local em 31/08/2026 e refletem a organização recente do menu superior e dos recursos integrados ao Tainacan.",
)

doc.add_heading("Menu superior e painel", level=2)
add_body_paragraph(
    doc,
    "O menu superior organiza o uso diário do sistema em Dashboard, Processos, Modelos, Itens do acervo, Tainacan, Configurações e WordPress. O Dashboard apresenta indicadores gerais e serve como ponto de partida para acompanhar a atividade do ambiente.",
)
add_bullets(
    doc,
    [
        "Dashboard: resumo do ambiente, indicadores e notificações.",
        "Processos: criação, acompanhamento, histórico, comentários e relatórios.",
        "Modelos: criação de modelos, edição de etapas e configuração de exportação.",
        "Itens do acervo: consulta aos itens do Tainacan e processos vinculados.",
        "Tainacan: acesso direto à administração do acervo.",
        "Configurações: acesso a Grupos e Usuários.",
        "WordPress: retorno ao painel administrativo geral.",
    ],
)
add_screenshot(doc, "dashboard.png", "Figura 26 - Dashboard atual do Obatalá.")

doc.add_heading("Processos", level=2)
add_body_paragraph(
    doc,
    "A tela Processos mostra o total de processos, as abas Todos os processos e Meus processos, filtros e a ação Adicionar novo. Em ambientes sem processos cadastrados, o sistema apresenta a mensagem de estado vazio, orientando que ainda não há registros.",
)
add_screenshot(doc, "processos.png", "Figura 27 - Listagem atual de processos.")
add_body_paragraph(
    doc,
    "Ao adicionar um processo, o usuário deve informar o título, escolher o modelo de processo e selecionar o nível de acesso. O nível Não restrito permite acesso mais amplo conforme as permissões do usuário; o nível Restrito limita a operação aos usuários autorizados pelo fluxo e pelos grupos associados.",
)
add_screenshot(doc, "processos-adicionar.png", "Figura 28 - Formulário atual para criação de processo.")

doc.add_heading("Modelos de processo e exportação Tainacan", level=2)
add_body_paragraph(
    doc,
    "A listagem de Modelos permite criar um novo modelo com título e descrição. Depois do cadastro, a ação principal para configuração é Gerenciar modelo. A antiga separação entre edição de etapas e edição dos dados de exportação foi unificada: o editor do modelo agora concentra o fluxo, os campos das etapas e o acordeon Exportação para o Tainacan.",
)
add_screenshot(doc, "modelos.png", "Figura 29 - Listagem atual de modelos.")
add_screenshot(doc, "modelos-adicionar.png", "Figura 30 - Cadastro atual de modelo de processo.")
add_body_paragraph(
    doc,
    "No editor do modelo, a configuração de exportação aparece no acordeon Exportação para o Tainacan, posicionado acima do acordeon Gerenciar etapas. Quando recolhido, ele exibe o estado atual do recurso, como Ativo ou Inativo. Ao expandi-lo, o usuário encontra o controle Exporter status. Depois de ativar esse controle, o sistema revela a área Coleções de destino e o seletor Selecione uma ou mais coleções.",
)
add_body_paragraph(
    doc,
    "Com a exportação habilitada e ao menos uma coleção selecionada, os componentes configurados dentro de cada etapa passam a exibir uma propriedade adicional de mapeamento. Essa propriedade aparece no bloco Tainacan Export das configurações do campo e permite ligar o campo do Obatalá a um metadado do Tainacan. Dessa forma, o dado preenchido durante o processo já fica preparado para compor o item exportado para a coleção escolhida.",
)
add_bullets(
    doc,
    [
        "Exporter status: ativa ou desativa o mapeamento para exportação.",
        "Coleções de destino: define as coleções do Tainacan que podem receber itens gerados pelo processo.",
        "Aviso de alerta: aparece quando a exportação está ativa, mas nenhuma coleção foi selecionada.",
        "Target collection: informa para qual coleção selecionada aquele campo será mapeado.",
        "Target metadata: informa qual metadado do Tainacan receberá o valor daquele campo.",
        "Do not map this field: mantém o campo apenas no formulário do Obatalá, sem envio para metadado do Tainacan.",
        "Resumo do mapeamento: mostra a relação Campo Obatalá x Metadado Tainacan por coleção.",
    ],
)
add_screenshot(doc, "tainacan-export-accordion-enabled.png", "Figura 31 - Acordeon Exportação para o Tainacan com o recurso ativado.")

doc.add_heading("Itens do acervo", level=2)
add_body_paragraph(
    doc,
    "A nova área Itens do acervo aproxima a gestão dos processos curatoriais dos registros mantidos no Tainacan. A listagem permite buscar e filtrar itens, visualizar coleção, situação, número de processos vinculados, última atualização e ações de consulta.",
)
add_screenshot(doc, "itens-acervo.png", "Figura 32 - Listagem de itens do acervo integrados ao Tainacan.")
add_body_paragraph(
    doc,
    "Ao abrir um item, a tela apresenta informações do registro, metadados, botões para abrir ou editar o item no Tainacan e a linha do tempo de processos vinculados. Essa visão ajuda a identificar rapidamente se um item já participou de processos curatoriais, se há processos concluídos, em progresso ou pendentes, e qual foi a última atualização.",
)
add_screenshot(doc, "item-detalhe-view.png", "Figura 33 - Detalhe de item do acervo e linha do tempo de processos.")

doc.add_heading("Grupos e usuários", level=2)
add_body_paragraph(
    doc,
    "A área Grupos lista os grupos cadastrados, seu status, descrição, quantidade de usuários e ações disponíveis. O detalhe do grupo mostra os usuários associados e permite administrar a composição do grupo de acordo com as permissões do usuário logado.",
)
add_screenshot(doc, "grupos.png", "Figura 34 - Listagem atual de grupos.")
add_screenshot(doc, "grupo-detalhe.png", "Figura 35 - Detalhe de grupo e usuários associados.")

doc.add_heading("Preparação e execução da exportação", level=2)
add_body_paragraph(
    doc,
    "Quando um modelo possui mapeamento Tainacan ativo, cada processo criado a partir dele apresenta uma etapa de controle chamada Preparação da exportação para o Tainacan. Nessa etapa, o usuário informa como os itens serão gerados: manualmente, preenchendo os dados no próprio processo, ou via planilha, enviando um arquivo estruturado com os valores que alimentarão os metadados mapeados.",
)
add_numbered(
    doc,
    [
        "Selecionar a coleção de destino quando houver mais de uma coleção mapeada.",
        "Definir se o processo gerará um item ou múltiplos itens no Tainacan.",
        "Para múltiplos itens, informar a quantidade e escolher a fonte dos dados: entrada manual ou planilha.",
        "Na entrada manual, preencher os valores diretamente nos campos exibidos nas etapas do processo.",
        "Na entrada por planilha, enviar arquivo CSV, XLS ou XLSX compatível com os campos mapeados.",
        "Opcionalmente repetir valores base entre itens e definir um campo identificador único com prefixo.",
        "Salvar a preparação da exportação antes de concluir a etapa de exportação.",
    ],
)
add_body_paragraph(
    doc,
    "A exportação cria ou atualiza itens no Tainacan com base nos metadados mapeados. O sistema também registra a referência do processo nos itens vinculados, facilitando a rastreabilidade entre o processo curatorial e o acervo.",
)

doc.add_heading("Relatórios e documentos", level=2)
add_body_paragraph(
    doc,
    "A listagem de processos pode exibir a ação Gerar relatório PDF. O relatório reúne dados do processo, número do processo, etapa atual, nível de acesso e informações registradas ao longo do fluxo. Campos do tipo Documento de etapa também podem gerar PDF e exigir anexação de PDF assinado quando essa opção estiver configurada no modelo.",
)

doc.add_heading("Permissões", level=2)
add_body_paragraph(
    doc,
    "A interface exibe menus e controles conforme as permissões da conta conectada. Administradores possuem acesso amplo; gestores de modelos podem configurar modelos; gestores de grupos administram grupos e usuários; gestores de mapeadores acessam os controles Tainacan Export; usuários de setor interagem com processos e etapas aos quais têm acesso. Em caso de ausência de permissão, a ação pode não aparecer ou retornar uma mensagem de acesso negado.",
)

doc.save(DOCX)
print(DOCX)
